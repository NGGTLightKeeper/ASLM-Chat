# Copyright NEXTGGTECH. Elastic License 2.0.

from __future__ import annotations

import base64
import html
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
from copy import deepcopy
from functools import lru_cache
from pathlib import Path
from typing import Any, Iterable


PROJECT_ROOT = Path(__file__).resolve().parents[2]
STATIC_ROOT = PROJECT_ROOT / "Apps" / "UI" / "static"
MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)", flags=re.I)
FENCED_BLOCK_RE = re.compile(r"(?ms)^(```|~~~)[^\n]*\n.*?^\1[ \t]*$")
DISPLAY_DOLLAR_RE = re.compile(r"(?s)(?<!\\)\$\$(.+?)(?<!\\)\$\$")
DISPLAY_BRACKET_RE = re.compile(r"(?s)\\\[(.+?)\\\]")
INLINE_PAREN_RE = re.compile(r"\\\((.+?)\\\)")
INLINE_DOLLAR_RE = re.compile(r"(?<!\\)\$(?!\s|\$)([^\n$]*?\S)(?<!\\)\$(?!\$)")


def safe_filename(title: str, suffix: str) -> str:
    stem = re.sub(r"[^\w.-]+", "-", str(title or "research"), flags=re.UNICODE).strip("-._")
    return f"{(stem or 'research')[:96]}.{suffix}"


def markdown_bytes(report: str) -> bytes:
    return str(report or "").strip().encode("utf-8")


def _encoded_latex_tag(source: str, *, display: bool) -> str:
    encoded = base64.b64encode(str(source or "").strip().encode("utf-8")).decode("ascii")
    kind = "display" if display else "inline"
    tag = "div" if display else "span"
    return f'<{tag} class="research-math research-math-{kind}" data-latex="{encoded}"></{tag}>'


def _protect_latex(markdown_text: str) -> str:
    """Replace LaTeX delimiters with Markdown-safe semantic placeholders."""

    protected: list[str] = []

    def hold_fence(match: re.Match[str]) -> str:
        protected.append(match.group(0))
        return f"\n\nASLMFENCEDBLOCK{len(protected) - 1}PLACEHOLDER\n\n"

    text = FENCED_BLOCK_RE.sub(hold_fence, str(markdown_text or ""))
    text = DISPLAY_DOLLAR_RE.sub(lambda match: _encoded_latex_tag(match.group(1), display=True), text)
    text = DISPLAY_BRACKET_RE.sub(lambda match: _encoded_latex_tag(match.group(1), display=True), text)
    text = INLINE_PAREN_RE.sub(lambda match: _encoded_latex_tag(match.group(1), display=False), text)
    text = INLINE_DOLLAR_RE.sub(lambda match: _encoded_latex_tag(match.group(1), display=False), text)
    for index, block in enumerate(protected):
        text = text.replace(f"ASLMFENCEDBLOCK{index}PLACEHOLDER", block)
    return text


def _decode_latex(value: Any) -> str:
    try:
        return base64.b64decode(str(value or ""), validate=True).decode("utf-8")
    except Exception:
        return ""


def _report_without_duplicate_title(report: str, title: str) -> str:
    text = str(report or "").strip()
    match = re.match(r"^#\s+(.+?)\s*(?:\n+|$)", text)
    if match and re.sub(r"\s+", " ", match.group(1)).casefold() == re.sub(r"\s+", " ", title).casefold():
        return text[match.end():].lstrip()
    return text


def _markdown_to_html(report: str) -> str:
    try:
        import markdown
    except Exception as exc:  # pragma: no cover - declared server dependency.
        raise RuntimeError("Markdown export rendering is unavailable.") from exc
    return markdown.markdown(
        _protect_latex(report),
        extensions=["extra", "sane_lists"],
        output_format="html5",
    )


def _chromium_candidates() -> Iterable[Path]:
    configured = str(os.environ.get("ASLM_CHROMIUM_PATH") or "").strip()
    if configured:
        yield Path(configured)
    for command in ("msedge", "chrome", "google-chrome", "chromium", "chromium-browser"):
        discovered = shutil.which(command)
        if discovered:
            yield Path(discovered)
    for environment_name, relative in (
        ("PROGRAMFILES(X86)", Path("Microsoft/Edge/Application/msedge.exe")),
        ("PROGRAMFILES", Path("Microsoft/Edge/Application/msedge.exe")),
        ("PROGRAMFILES", Path("Google/Chrome/Application/chrome.exe")),
        ("PROGRAMFILES(X86)", Path("Google/Chrome/Application/chrome.exe")),
    ):
        root = str(os.environ.get(environment_name) or "").strip()
        if root:
            yield Path(root) / relative


@lru_cache(maxsize=1)
def _chromium_path() -> Path:
    for candidate in _chromium_candidates():
        if candidate.is_file():
            return candidate.resolve()
    raise RuntimeError("PDF and diagram export require Microsoft Edge, Chrome, or Chromium.")


def _run_chromium(arguments: list[str], *, timeout: float = 45.0) -> None:
    command = [str(_chromium_path()), *arguments]
    completed = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
        check=False,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    if completed.returncode != 0:
        detail = completed.stderr.decode("utf-8", errors="replace").strip()[-1200:]
        raise RuntimeError(f"Chromium export failed: {detail or completed.returncode}")


def _trim_screenshot(path: Path, *, padding: int = 22) -> None:
    try:
        from PIL import Image, ImageChops
    except Exception:  # pragma: no cover - Pillow is already a server dependency.
        return
    with Image.open(path) as source:
        # Chromium on Windows uses RGB subpixel antialiasing even for black
        # text. Collapse the channels so exported research is truly
        # monochrome rather than merely styled with black CSS tokens.
        image = source.convert("L").convert("RGB")
        background = Image.new("RGB", image.size, image.getpixel((0, 0)))
        difference = ImageChops.difference(image, background)
        box = difference.getbbox()
        if not box:
            return
        left = max(0, box[0] - padding)
        top = max(0, box[1] - padding)
        right = min(image.width, box[2] + padding)
        bottom = min(image.height, box[3] + padding)
        image.crop((left, top, right, bottom)).save(path, "PNG")


def _browser_asset(path: str) -> str:
    asset = STATIC_ROOT / Path(path)
    if not asset.is_file():
        raise RuntimeError(f"Research export asset is missing: {asset.name}")
    return asset.resolve().as_uri()


def _mermaid_monochrome_config() -> str:
    return json.dumps({
        "startOnLoad": False,
        "securityLevel": "strict",
        "theme": "base",
        "themeVariables": {
            "background": "#ffffff",
            "primaryColor": "#ffffff",
            "primaryTextColor": "#000000",
            "primaryBorderColor": "#000000",
            "lineColor": "#000000",
            "secondaryColor": "#ffffff",
            "secondaryTextColor": "#000000",
            "secondaryBorderColor": "#000000",
            "tertiaryColor": "#ffffff",
            "tertiaryTextColor": "#000000",
            "tertiaryBorderColor": "#000000",
            "noteBkgColor": "#ffffff",
            "noteTextColor": "#000000",
            "noteBorderColor": "#000000",
            "clusterBkg": "#ffffff",
            "clusterBorder": "#000000",
            "actorBkg": "#ffffff",
            "actorBorder": "#000000",
            "actorTextColor": "#000000",
            "signalColor": "#000000",
            "signalTextColor": "#000000",
            "labelBoxBkgColor": "#ffffff",
            "labelBoxBorderColor": "#000000",
            "labelTextColor": "#000000",
            "loopTextColor": "#000000",
            "activationBkgColor": "#ffffff",
            "activationBorderColor": "#000000",
            "fontFamily": "Arial, sans-serif",
        },
    }, ensure_ascii=True)


def _render_mermaid_png(source: str, output_path: Path, workspace: Path) -> None:
    encoded_source = json.dumps(str(source or ""), ensure_ascii=False).replace("<", "\\u003c")
    mermaid_config = _mermaid_monochrome_config()
    document = f"""<!doctype html>
<html><head><meta charset="utf-8"><style>
html,body{{margin:0;background:#fff}}body{{display:inline-block;padding:24px;font-family:Arial,sans-serif}}
#diagram svg{{display:block;max-width:1450px;height:auto;background:#fff!important;color:#000!important}}
#diagram svg *{{color:#000!important}}
#diagram svg text,#diagram svg tspan{{fill:#000!important}}
#diagram svg rect,#diagram svg polygon,#diagram svg circle,#diagram svg ellipse{{fill:#fff!important;stroke:#000!important}}
#diagram svg path,#diagram svg line,#diagram svg polyline{{stroke:#000!important}}
#diagram svg marker path{{fill:#000!important;stroke:#000!important}}
#diagram svg .edgeLabel,#diagram svg .label,#diagram svg .label-container{{background:#fff!important;color:#000!important}}
</style></head><body><div id="diagram"></div>
<script src="{_browser_asset('js/vendor/mermaid.min.js')}"></script>
<script>(async()=>{{
  mermaid.initialize({mermaid_config});
  const result=await mermaid.render('aslm-export-diagram',{encoded_source});
  document.getElementById('diagram').innerHTML=result.svg;
  document.title='READY';
}})().catch(error=>{{document.body.textContent=String(error);document.title='ERROR'}});</script>
</body></html>"""
    html_path = workspace / f"{output_path.stem}.html"
    html_path.write_text(document, encoding="utf-8")
    profile = workspace / f"profile-{output_path.stem}"
    _run_chromium([
        "--headless=new",
        "--disable-gpu",
        "--disable-extensions",
        "--hide-scrollbars",
        "--allow-file-access-from-files",
        "--run-all-compositor-stages-before-draw",
        "--virtual-time-budget=10000",
        "--force-device-scale-factor=3",
        "--window-size=1600,1000",
        f"--user-data-dir={profile}",
        f"--screenshot={output_path}",
        html_path.resolve().as_uri(),
    ])
    if not output_path.is_file() or output_path.stat().st_size < 500:
        raise RuntimeError("Mermaid diagram did not render to an image.")
    _trim_screenshot(output_path)


def _render_latex_png(source: str, output_path: Path, workspace: Path) -> None:
    encoded_source = json.dumps(str(source or ""), ensure_ascii=False).replace("<", "\\u003c")
    document = f"""<!doctype html>
<html><head><meta charset="utf-8"><link rel="stylesheet" href="{_browser_asset('css/vendor/katex.min.css')}">
<style>html,body{{margin:0;background:#fff}}body{{display:inline-block;padding:20px;color:#000}}</style></head>
<body><div id="formula"></div><script src="{_browser_asset('js/vendor/katex.min.js')}"></script>
<script>katex.render({encoded_source},document.getElementById('formula'),{{displayMode:true,throwOnError:false,strict:false}});document.title='READY';</script>
</body></html>"""
    html_path = workspace / f"{output_path.stem}.html"
    html_path.write_text(document, encoding="utf-8")
    profile = workspace / f"profile-{output_path.stem}"
    _run_chromium([
        "--headless=new",
        "--disable-gpu",
        "--hide-scrollbars",
        "--allow-file-access-from-files",
        "--virtual-time-budget=4000",
        "--force-device-scale-factor=2",
        "--window-size=1400,320",
        f"--user-data-dir={profile}",
        f"--screenshot={output_path}",
        html_path.resolve().as_uri(),
    ])
    if not output_path.is_file() or output_path.stat().st_size < 300:
        raise RuntimeError("LaTeX formula did not render to an image.")
    _trim_screenshot(output_path, padding=12)


def _mml2omml_path() -> Path | None:
    configured = str(os.environ.get("ASLM_MML2OMML_XSL") or "").strip()
    candidates = [Path(configured)] if configured else []
    for environment_name in ("PROGRAMFILES", "PROGRAMFILES(X86)"):
        root = str(os.environ.get(environment_name) or "").strip()
        if root:
            candidates.append(Path(root) / "Microsoft Office" / "root" / "Office16" / "MML2OMML.XSL")
    return next((candidate.resolve() for candidate in candidates if candidate.is_file()), None)


@lru_cache(maxsize=1)
def _math_transform():
    path = _mml2omml_path()
    if path is None:
        return None
    from lxml import etree
    return etree.XSLT(etree.parse(str(path)))


def _latex_omml(source: str):
    transform = _math_transform()
    if transform is None:
        return None
    try:
        from latex2mathml.converter import convert
        from lxml import etree
        mathml = convert(str(source or "").strip())
        transformed = transform(etree.fromstring(mathml.encode("utf-8")))
        return deepcopy(transformed.getroot())
    except Exception:
        return None


def _append_latex(paragraph, source: str, *, display: bool, workspace: Path, image_index: list[int]) -> None:
    omml = _latex_omml(source)
    if omml is not None:
        paragraph._p.append(omml)
        return
    from docx.shared import Inches
    image_index[0] += 1
    output = workspace / f"formula-{image_index[0]}.png"
    _render_latex_png(source, output, workspace)
    try:
        from PIL import Image
        with Image.open(output) as rendered:
            width = min(6.2 if display else 2.8, max(0.35, rendered.width / 160.0))
    except Exception:
        width = 2.8 if display else 1.5
    paragraph.add_run().add_picture(str(output), width=Inches(width))


def _set_run_font(run, name: str) -> None:
    from docx.oxml.ns import qn
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)


def _append_text_run(paragraph, text: str, formatting: dict[str, bool]) -> None:
    from docx.shared import RGBColor
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bool(formatting.get("bold"))
    run.italic = bool(formatting.get("italic"))
    if formatting.get("code"):
        _set_run_font(run, "Consolas")


def _append_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    relationship_id = paragraph.part.relate_to(str(url), RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = str(text or url)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_inline(paragraph, node, workspace: Path, image_index: list[int], formatting: dict[str, bool] | None = None) -> None:
    from bs4 import NavigableString, Tag
    style = dict(formatting or {})
    if isinstance(node, NavigableString):
        _append_text_run(paragraph, str(node), style)
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    if node.has_attr("data-latex"):
        source = _decode_latex(node.get("data-latex"))
        if source:
            _append_latex(paragraph, source, display=name == "div", workspace=workspace, image_index=image_index)
        return
    if name == "br":
        paragraph.add_run().add_break()
        return
    if name == "a":
        _append_hyperlink(paragraph, node.get_text(" ", strip=False), str(node.get("href") or ""))
        return
    if name in {"strong", "b"}:
        style["bold"] = True
    elif name in {"em", "i"}:
        style["italic"] = True
    elif name == "code":
        style["code"] = True
    for child in node.children:
        _append_inline(paragraph, child, workspace, image_index, style)


def _shade(element, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    properties = element.get_or_add_pPr() if hasattr(element, "get_or_add_pPr") else element.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_table_geometry(table, widths: list[int]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    table.autofit = False
    properties = table._tbl.tblPr
    width_node = properties.find(qn("w:tblW"))
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        properties.append(width_node)
    width_node.set(qn("w:type"), "dxa")
    width_node.set(qn("w:w"), str(sum(widths)))
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(widths[index]))


def _table_widths(rows: list[list[Any]]) -> list[int]:
    column_count = max((len(row) for row in rows), default=1)
    weights = []
    for index in range(column_count):
        longest = max((len(row[index].get_text(" ", strip=True)) if index < len(row) else 0 for row in rows), default=1)
        weights.append(max(9.0, min(34.0, longest ** 0.62)))
    total = sum(weights) or 1.0
    widths = [max(900, round(9360 * weight / total)) for weight in weights]
    widths[-1] += 9360 - sum(widths)
    return widths


def _append_table(document, tag, workspace: Path, image_index: list[int]) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    rows = [row.find_all(["th", "td"], recursive=False) for row in tag.find_all("tr")]
    rows = [row for row in rows if row]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    widths = _table_widths(rows)
    for row_index, source_row in enumerate(rows):
        for column_index, source_cell in enumerate(source_row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = 0
            for child in source_cell.children:
                _append_inline(paragraph, child, workspace, image_index)
            if row_index == 0 and source_cell.name.lower() == "th":
                _shade(cell._tc, "FFFFFF")
                for run in paragraph.runs:
                    run.bold = True
    _set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = 0


def _append_code_block(document, source: str) -> None:
    from docx.shared import Pt
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    _shade(paragraph._p, "FFFFFF")
    run = paragraph.add_run(str(source or "").rstrip())
    _set_run_font(run, "Consolas")
    run.font.size = Pt(9)


def _append_mermaid(document, source: str, workspace: Path, image_index: list[int]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches
    image_index[0] += 1
    output = workspace / f"mermaid-{image_index[0]}.png"
    try:
        _render_mermaid_png(source, output, workspace)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        try:
            from PIL import Image
            with Image.open(output) as rendered:
                width = min(6.35, max(5.35, rendered.width / 130.0))
        except Exception:
            width = 5.75
        paragraph.add_run().add_picture(str(output), width=Inches(width))
    except Exception:
        _append_code_block(document, source)


def _append_list(document, tag, workspace: Path, image_index: list[int], *, ordered: bool, level: int = 0) -> None:
    from docx.shared import Inches
    for item in tag.find_all("li", recursive=False):
        paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5 + (level * 0.25))
        for child in item.children:
            if getattr(child, "name", "") in {"ul", "ol"}:
                continue
            _append_inline(paragraph, child, workspace, image_index)
        for nested in item.find_all(["ul", "ol"], recursive=False):
            _append_list(document, nested, workspace, image_index, ordered=nested.name.lower() == "ol", level=level + 1)


def _append_blocks(document, container, workspace: Path, image_index: list[int]) -> None:
    from bs4 import NavigableString, Tag
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    for child in container.children:
        if isinstance(child, NavigableString):
            if str(child).strip():
                document.add_paragraph(str(child).strip())
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name.lower()
        if child.has_attr("data-latex"):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _append_latex(paragraph, _decode_latex(child.get("data-latex")), display=True, workspace=workspace, image_index=image_index)
        elif re.fullmatch(r"h[1-6]", name):
            paragraph = document.add_heading(level=min(3, int(name[1])))
            for inline in child.children:
                _append_inline(paragraph, inline, workspace, image_index)
        elif name == "p":
            paragraph = document.add_paragraph()
            for inline in child.children:
                _append_inline(paragraph, inline, workspace, image_index)
        elif name in {"ul", "ol"}:
            _append_list(document, child, workspace, image_index, ordered=name == "ol")
        elif name == "table":
            _append_table(document, child, workspace, image_index)
        elif name == "pre":
            code = child.find("code")
            classes = {str(value).lower() for value in (code.get("class") if code else [])}
            source = code.get_text() if code else child.get_text()
            if "language-mermaid" in classes:
                _append_mermaid(document, source, workspace, image_index)
            else:
                _append_code_block(document, source)
        elif name == "blockquote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.28)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(8)
            _shade(paragraph._p, "FFFFFF")
            for inline in child.descendants:
                if isinstance(inline, NavigableString) and str(inline).strip():
                    _append_text_run(paragraph, str(inline), {"italic": True})
        elif name == "hr":
            paragraph = document.add_paragraph("\u2500" * 38)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            _append_blocks(document, child, workspace, image_index)


def _configure_document(document, title: str) -> None:
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.space_after = Pt(12)
    title_properties = title_style.element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)
    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    document.add_paragraph(str(title or "Research report"), style="Title")


def docx_bytes(report: str, title: str) -> bytes:
    try:
        from bs4 import BeautifulSoup
        from docx import Document
    except Exception as exc:  # pragma: no cover - declared server dependencies.
        raise RuntimeError("Word export dependencies are unavailable.") from exc
    normalized_title = str(title or "Research report").strip()
    body = _report_without_duplicate_title(report, normalized_title)
    soup = BeautifulSoup(_markdown_to_html(body), "html.parser")
    with tempfile.TemporaryDirectory(prefix="aslm-research-docx-") as directory:
        workspace = Path(directory)
        document = Document()
        _configure_document(document, normalized_title)
        _append_blocks(document, soup, workspace, [0])
        from io import BytesIO
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()


def _pdf_html(report: str, title: str) -> str:
    encoded_title = html.escape(str(title or "Research report"))
    prepared_report = _protect_latex(_report_without_duplicate_title(report, title))
    encoded_report = json.dumps(prepared_report, ensure_ascii=False).replace("<", "\\u003c")
    mermaid_config = _mermaid_monochrome_config()
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><title>Rendering report</title>
<link rel="stylesheet" href="{_browser_asset('css/vendor/katex.min.css')}">
<style>
@page{{size:Letter;margin:17mm 18mm 20mm}}*{{box-sizing:border-box}}
html,body{{margin:0;padding:0;background:#fff;color:#000;font-family:Arial,'Segoe UI',sans-serif;font-size:10.5pt;line-height:1.55}}
main{{width:100%}}h1.report-title{{margin:0 0 20pt;color:#000;font-size:23pt;line-height:1.15}}
h1,h2,h3,h4{{color:#000;line-height:1.24;break-after:avoid-page}}h1{{font-size:18pt;margin:22pt 0 9pt}}h2{{font-size:15pt;margin:18pt 0 7pt}}h3{{font-size:12.5pt;margin:14pt 0 5pt}}
p{{margin:0 0 8pt}}a{{color:#000;text-decoration:underline}}strong{{font-weight:700}}
ul,ol{{margin:4pt 0 10pt;padding-left:22pt}}li{{margin:2pt 0}}blockquote{{margin:10pt 0;padding:8pt 12pt;border-left:3px solid #000;background:#fff}}
table{{width:100%;border-collapse:collapse;margin:10pt 0 14pt;font-size:9.5pt;break-inside:auto}}thead{{display:table-header-group}}tr{{break-inside:avoid}}
th,td{{padding:6pt 7pt;border:0.6pt solid #000;vertical-align:top;text-align:left}}th{{background:#fff;font-weight:700}}
pre{{white-space:pre-wrap;overflow-wrap:anywhere;margin:9pt 0 12pt;padding:9pt 11pt;border:0.6pt solid #000;border-radius:4pt;background:#fff;font:8.8pt/1.42 Consolas,'Courier New',monospace;break-inside:avoid-page}}
code{{font-family:Consolas,'Courier New',monospace;font-size:.92em}}:not(pre)>code{{padding:1pt 3pt;border:1px solid #000;border-radius:3pt;background:#fff}}
.research-math-display{{margin:12pt 0;text-align:center;break-inside:avoid-page}}.research-math-inline{{white-space:nowrap}}
.mermaid-export{{margin:12pt auto 15pt;text-align:center;break-inside:avoid-page}}.mermaid-export svg{{max-width:100%;height:auto;background:#fff!important;color:#000!important}}
.mermaid-export svg *{{color:#000!important}}.mermaid-export svg text,.mermaid-export svg tspan{{fill:#000!important}}
.mermaid-export svg rect,.mermaid-export svg polygon,.mermaid-export svg circle,.mermaid-export svg ellipse{{fill:#fff!important;stroke:#000!important}}
.mermaid-export svg path,.mermaid-export svg line,.mermaid-export svg polyline{{stroke:#000!important}}.mermaid-export svg marker path{{fill:#000!important;stroke:#000!important}}
.mermaid-export svg .edgeLabel,.mermaid-export svg .label,.mermaid-export svg .label-container{{background:#fff!important;color:#000!important}}
hr{{border:0;border-top:.7pt solid #000;margin:16pt 0}}img,svg{{max-width:100%}}
</style></head><body><main><h1 class="report-title">{encoded_title}</h1><article id="report"></article></main>
<script src="{_browser_asset('js/vendor/marked.min.js')}"></script>
<script src="{_browser_asset('js/vendor/purify.min.js')}"></script>
<script src="{_browser_asset('js/vendor/katex.min.js')}"></script>
<script src="{_browser_asset('js/vendor/mermaid.min.js')}"></script>
<script>(async()=>{{
  const root=document.getElementById('report');
  root.innerHTML=DOMPurify.sanitize(marked.parse({encoded_report}));
  const decode=value=>new TextDecoder().decode(Uint8Array.from(atob(value),character=>character.charCodeAt(0)));
  root.querySelectorAll('[data-latex]').forEach(node=>{{
    katex.render(decode(node.dataset.latex),node,{{displayMode:node.classList.contains('research-math-display'),throwOnError:false,strict:false}});
  }});
  mermaid.initialize({mermaid_config});
  const diagrams=[...root.querySelectorAll('pre code.language-mermaid')];
  for(let index=0;index<diagrams.length;index+=1){{
    const code=diagrams[index];const host=document.createElement('div');host.className='mermaid-export';
    code.closest('pre').replaceWith(host);const rendered=await mermaid.render(`aslm-pdf-${{index}}`,code.textContent||'');host.innerHTML=rendered.svg;
  }}
  document.title='ASLM_EXPORT_READY';
}})().catch(error=>{{document.body.insertAdjacentHTML('beforeend','<pre>'+String(error)+'</pre>');document.title='ASLM_EXPORT_ERROR'}});</script>
</body></html>"""


def pdf_bytes(report: str, title: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="aslm-research-pdf-") as directory:
        workspace = Path(directory)
        html_path = workspace / "report.html"
        pdf_path = workspace / "report.pdf"
        profile = workspace / "profile"
        html_path.write_text(_pdf_html(report, title), encoding="utf-8")
        _run_chromium([
            "--headless=new",
            "--disable-gpu",
            "--disable-extensions",
            "--allow-file-access-from-files",
            "--run-all-compositor-stages-before-draw",
            "--virtual-time-budget=15000",
            "--print-to-pdf-no-header",
            "--no-pdf-header-footer",
            f"--user-data-dir={profile}",
            f"--print-to-pdf={pdf_path}",
            html_path.resolve().as_uri(),
        ], timeout=60.0)
        deadline = time.monotonic() + 5.0
        while time.monotonic() < deadline and (not pdf_path.is_file() or pdf_path.stat().st_size < 1000):
            time.sleep(0.05)
        if not pdf_path.is_file():
            raise RuntimeError("Chromium did not produce the research PDF.")
        payload = pdf_path.read_bytes()
        if not payload.startswith(b"%PDF-"):
            raise RuntimeError("Chromium produced an invalid research PDF.")
        return payload
